"""
Document ingestion and indexing module for RAG AI Decision Assistant
Supports PDF, DOCX, and text files in Russian and English
"""
import os
import logging
import pickle
from pathlib import Path
from typing import List
import pdfplumber
from docx import Document

from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings

from config import settings

# Try to import sentence transformers for local embeddings
try:
    from langchain_community.embeddings import HuggingFaceEmbeddings
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# Configure logging
logging.basicConfig(
    level=getattr(logging, settings.log_level),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def load_text(file_path: str) -> str:
    """
    Load text from various file formats (PDF, DOCX, TXT)
    Supports Russian and English text
    
    Args:
        file_path: Path to the file to load
        
    Returns:
        Extracted text content
        
    Raises:
        Exception: If file cannot be read or processed
    """
    text = ""
    file_path_obj = Path(file_path)
    
    try:
        if file_path.endswith(".pdf"):
            logger.info(f"Loading PDF: {file_path}")
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif file_path.endswith(".docx"):
            logger.info(f"Loading DOCX: {file_path}")
            doc = Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
        elif file_path.endswith((".txt", ".md")):
            logger.info(f"Loading text file: {file_path}")
            # Try UTF-8 first, fallback to cp1252 (common on Windows) or ignore errors
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decode failed for {file_path}, falling back to cp1252")
                with open(file_path, "r", encoding="cp1252", errors="replace") as f:
                    text = f.read()
        else:
            logger.warning(f"Unsupported file type: {file_path}")
            return ""
            
        if not text.strip():
            logger.warning(f"No text extracted from {file_path}")
            
        return text
        
    except Exception as e:
        logger.error(f"Error loading {file_path}: {str(e)}")
        raise


def ingest_knowledge_base() -> None:
    """
    Main ingestion function: loads documents, chunks them, creates embeddings,
    and saves the FAISS index
    
    Raises:
        Exception: If ingestion fails at any step
    """
    data_dir = Path(settings.data_dir)
    
    # Validate data directory exists
    if not data_dir.exists():
        raise FileNotFoundError(
            f"Data directory '{settings.data_dir}' not found. "
            "Please create it and add your knowledge base documents."
        )
    
    # Get all supported files
    supported_extensions = {".pdf", ".docx", ".txt", ".md"}
    files = [
        f for f in data_dir.iterdir()
        if f.is_file() and f.suffix.lower() in supported_extensions
    ]
    
    if not files:
        raise ValueError(
            f"No supported files found in '{settings.data_dir}'. "
            f"Supported formats: PDF, DOCX, TXT, MD"
        )
    
    logger.info(f"Found {len(files)} files to process")
    
    # Load all documents
    all_docs = []
    for file_path in files:
        try:
            text = load_text(str(file_path))
            if text.strip():
                # Create LangChain Document objects with source metadata
                all_docs.append(LCDocument(page_content=text, metadata={"source": file_path.name}))
                logger.info(f"Successfully loaded {file_path.name} ({len(text)} characters)")
        except Exception as e:
            logger.error(f"Failed to load {file_path.name}: {str(e)}")
            continue
    
    if not all_docs:
        raise ValueError("No text content could be extracted from any files")
    
    # Chunk documents
    logger.info("Chunking documents...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        length_function=len
    )
    
    docs = text_splitter.split_documents(all_docs)
    
    logger.info(f"Created {len(docs)} text chunks")
    
    if not docs:
        raise ValueError("No text chunks created from documents")
    
    # Create embeddings and vector store
    logger.info("Creating embeddings and vector store...")
    try:
        embeddings = None
        use_local_embeddings = False
        
        # Try OpenAI embeddings first
        if settings.openai_api_key and settings.openai_api_key != "your_openai_api_key_here":
            try:
                logger.info("Attempting to use OpenAI embeddings...")
                embeddings = OpenAIEmbeddings(
                    model=settings.embedding_model,
                    openai_api_key=settings.openai_api_key
                )
                # Test the embeddings with a small sample
                test_emb = embeddings.embed_query("test")
                logger.info("OpenAI embeddings working successfully")
            except Exception as e:
                if "429" in str(e) or "quota" in str(e).lower():
                    logger.warning(f"OpenAI quota exceeded: {str(e)}")
                    logger.info("Falling back to local embeddings...")
                    use_local_embeddings = True
                else:
                    raise
        
        # Use local embeddings if OpenAI failed or not available
        if embeddings is None or use_local_embeddings:
            if SENTENCE_TRANSFORMERS_AVAILABLE:
                logger.info("Using local HuggingFace embeddings (sentence-transformers)...")
                # Use a multilingual model that supports both English and Russian
                embeddings = HuggingFaceEmbeddings(
                    model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
                )
            else:
                raise ValueError(
                    "No embedding method available. Either:\n"
                    "1. Set a valid OPENAI_API_KEY with available quota, or\n"
                    "2. Install sentence-transformers: pip install sentence-transformers"
                )
        
        vector_store = FAISS.from_documents(docs, embeddings)
        logger.info("Vector store created successfully")
    except Exception as e:
        logger.error(f"Error creating embeddings: {str(e)}")
        raise
    
    # Save FAISS index
    logger.info(f"Saving index to {settings.index_path}...")
    import shutil
    import time
    
    temp_index_path = f"{settings.index_path}_temp"
    try:
        # Create parent dir for index
        Path(settings.index_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Save to a temporary location first
        if os.path.exists(temp_index_path):
            shutil.rmtree(temp_index_path)
        vector_store.save_local(temp_index_path)
        
        # Now try to replace the old one with retries for Windows locks
        max_retries = 5
        for i in range(max_retries):
            try:
                if os.path.exists(settings.index_path):
                    shutil.rmtree(settings.index_path)
                shutil.move(temp_index_path, settings.index_path)
                logger.info(f"Index saved successfully to {settings.index_path}")
                break
            except PermissionError as e:
                if i < max_retries - 1:
                    logger.warning(f"File locked, retrying in 1s... ({i+1}/{max_retries})")
                    time.sleep(1)
                else:
                    raise e
    except Exception as e:
        logger.error(f"Error saving index: {str(e)}")
        if os.path.exists(temp_index_path):
            try:
                shutil.rmtree(temp_index_path)
            except:
                pass
        raise
    
    logger.info("Knowledge base ingested and indexed successfully")


def fetch_url_text(url: str) -> str:
    """
    Fetch a web page and extract its plain text.
    Uses realistic browser headers to avoid bot-detection blocks (403/429).

    Args:
        url: Full HTTP/HTTPS URL to fetch

    Returns:
        Cleaned plain text with title and URL header

    Raises:
        ValueError: If the site blocks the request or returns no usable content
        Exception: On network or parsing failure
    """
    import re
    import requests
    from bs4 import BeautifulSoup

    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.5",
        "DNT": "1",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
    }

    session = requests.Session()
    resp = session.get(url, headers=headers, timeout=20, allow_redirects=True)

    if resp.status_code == 403:
        raise ValueError(
            f"The website at {url} blocked the request (HTTP 403). "
            "This site likely requires a real browser or login. "
            "Try saving the page content manually as a .txt or .pdf file and uploading it instead."
        )
    if resp.status_code == 429:
        raise ValueError(
            f"Rate limited by {url} (HTTP 429). Wait a moment and try again."
        )
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "lxml")

    # Drop non-content tags
    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript", "form"]):
        tag.decompose()

    title = soup.title.get_text(strip=True) if soup.title else ""
    body_text = soup.get_text(separator="\n", strip=True)
    body_text = re.sub(r"\n{3,}", "\n\n", body_text)

    if len(body_text.strip()) < 100:
        raise ValueError(
            f"Could not extract meaningful text from {url}. "
            "The page may require JavaScript or be behind a login wall. "
            "Try saving the content as a file and uploading it instead."
        )

    prefix = f"Title: {title}\nURL: {url}\n\n" if title else f"URL: {url}\n\n"
    return prefix + body_text


if __name__ == "__main__":
    try:
        ingest_knowledge_base()
        print("Knowledge base ingestion completed successfully!")
    except Exception as e:
        logger.error(f"Ingestion failed: {str(e)}")
        exit(1)
